import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMP_DIR = "temp_files"  # Temporary directory path

def CO_Table():
    # List all .docx files in the TEMP_DIR
    file_names = [f for f in os.listdir(TEMP_DIR) if f.startswith("18") and f.endswith(".docx")]
    # Create a new document to accumulate content from all files
    new_doc = Document('output.docx')

    for file_path in file_names:
        # Open the current document
        print(file_path)
        doc = Document(TEMP_DIR + "\\" + file_path)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Ignore empty paragraphs
                if "Course Code and Name:" in paragraph.text:
                    CourseDetails = paragraph.text.split("Course Code and Name:")[1].strip()
                    if "–" in CourseDetails:
                        course_code = CourseDetails.split("–")[0].strip()
                        course_name = CourseDetails.split("–")[1].strip()
                if "Class and Semester:" in paragraph.text:
                    ClassDetails = paragraph.text.split("Class and Semester:")[1].strip()

        para = new_doc.add_paragraph(f"Class: {ClassDetails} Subject: {course_name}", style=paragraph.style)
        # Make the text bold
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(14)  # Adjust the size to your preference, e.g., 14 points
        # Center-align the paragraph
        para.alignment = 1  # 1 corresponds to center alignment






        # Add a table for course outcomes
        new_table = new_doc.add_table(rows=0, cols=5)
        new_table.style = 'Table Grid'
        # Add a header row to the new table
        header_row = new_table.add_row()
        header_row.cells[0].text = "Course Code"
        header_row.cells[1].text = "Course Outcome"
        # Center-align the text in the header cells
        header_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Iterate through tables in the document
        for table in doc.tables:
            if "Course Outcomes" in table.cell(0, 0).text:  # Check for "Course Outcomes" in the first cell
                if course_code in table.cell(0, 1).text:  # Check for "MEC 701" in the second cell
                    cell_text = table.cell(0, 1).text  # Get the text from the cell
                    for i in range(1, 7):  # Loop through patterns MEC 701.1 to MEC 701.6
                        pattern = f"{course_code}.{i}"
                        if pattern in cell_text:
                            # Extract text associated with the current pattern
                            parts = cell_text.split(pattern, maxsplit=1)
                            description = parts[1].split(f"{course_code}.{i + 1}")[0].strip() if f"{course_code}.{i + 1}" in parts[1] else parts[1].strip()

                            # Add a new row to the new table
                            new_row = new_table.add_row()
                            new_row.cells[0].text = pattern  # Add the pattern to the first column
                            new_row.cells[1].text = description  # Add the description to the second column

                            # Center-align the pattern column (first column)
                            new_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Merge cells in columns 2, 3, and 4 with the first column if blank
        for row in range(len(new_table.rows)):
            for col in range(3, 5): 
                top_cell = new_table.cell(row, col)
                first_column_cell = new_table.cell(row, 1)

                if not top_cell.text.strip() and first_column_cell.text.strip():
                    top_cell.merge(first_column_cell)  # Merge with the first column

        # Add space below the table (empty paragraphs)
        new_doc.add_paragraph("")  # First blank line below the table
        new_doc.add_paragraph("")  # Second blank line below the table

    # Save everything into one document
    output_file = 'output.docx'
    new_doc.save(output_file)
    os.startfile(output_file)
