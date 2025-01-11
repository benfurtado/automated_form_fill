from docx import Document
import os

def fill_word_form(template_path, output_path, data):
    try:
        # Check if the template file exists
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        # Load the Word document
        doc = Document(template_path)
        
        # Replace placeholders in text
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        # Save the updated document
        doc.save(output_path)
        print(f"Form successfully filled and saved to {output_path}")

    except FileNotFoundError as e:
        print(e)
    except Exception as e:
        print(f"An error occurred: {e}")

# Example usage
template = "form_template.docx"
output = "filled_form.docx"
form_data = {
    "{{Name}}": "Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to Form successfully filled and saved to ",
    "{{Date}}": "2024-12-30",
    "{{Address}}": "123 Main Street, Springfield",
}

fill_word_form(template, output, form_data)
