import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

# Load workbook and sheet
workbook = load_workbook('14 15 PO and PSO.xlsx')
sheet = workbook['Sheet1']

# Open the existing Word document
doc = Document('Excel_to_Word_Table.docx')

# Variables for tracking current heading and table
current_table = None

# Process each row in the Excel sheet
for row in sheet.iter_rows(values_only=True):
    skip_row = False
    for cell in row:
        if cell and isinstance(cell, str):
            # Parse "heading:" and "text:"
            heading = None
            text = None
            if "heading:" in cell.lower():
                heading = cell.split("heading:")[1].split("text:")[0].strip() if "text:" in cell.lower() else cell.split("heading:")[1].strip()
            if "text:" in cell.lower():
                text = cell.split("text:")[1].strip() if "text:" in cell.lower() else None

            if heading or text:
                if heading:
                    doc.add_paragraph(heading, style='Heading 1')
                if text:
                    doc.add_paragraph(text, style='Normal')

                # Create a new table with 6 columns
                current_table = doc.add_table(rows=0, cols=6)
                current_table.style = 'Table Grid'
                skip_row = True
                break

    if skip_row:
        continue  # Skip rows with "heading:" or "text:"

    # Add data to the table if it exists
    if current_table:
        # Filter non-empty cells and limit to 6 columns
        non_empty_row = [str(cell).strip() for cell in row if cell]  # Convert all valid cells to strings
        if non_empty_row:
            word_row = current_table.add_row().cells[:6]  # Ensure no extra columns

            # Fill the row
            for col_index, value in enumerate(non_empty_row[:6]):
                cell = word_row[col_index]
                run = cell.paragraphs[0].add_run(value)
                if col_index == 0:  # Bold the first column
                    run.bold = True
                run.font.size = Pt(10.5)

            # Merge the last 4 columns
            if len(word_row) > 1:
                word_row[1].merge(word_row[2])
                word_row[1].merge(word_row[3])
                word_row[1].merge(word_row[4])
                word_row[1].merge(word_row[5])

# Save and open the Word document
word_file_path = 'Excel_to_Word_Table.docx'
doc.save(word_file_path)
print("Data written to Word table successfully and saved!")
os.startfile(word_file_path)
