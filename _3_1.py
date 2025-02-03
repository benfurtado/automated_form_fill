import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

TEMP_DIR = "temp_files"
OUTPUT_FILE = "Output.docx"

def PO_PSO():
    file_names = [f for f in os.listdir(TEMP_DIR) if f.startswith("14") and f.endswith(".xlsx")]

    # Load existing document or create a new one
    if os.path.exists(OUTPUT_FILE):
        new_doc = Document(OUTPUT_FILE)
    else:
        new_doc = Document()

    for file_name in file_names:
        file_path = os.path.join(TEMP_DIR, file_name)
        print(f"Processing: {file_name}")

        try:
            # Load the workbook
            workbook = load_workbook(file_path)
            sheet = workbook.active  # Automatically selects the first sheet

            current_table = None  # Variable for tracking the table

            for row in sheet.iter_rows(values_only=True):
                skip_row = False
                heading, text = None, None

                # Extract heading and text if present
                for cell in row:
                    if cell and isinstance(cell, str):
                        cell_lower = cell.lower()
                        if "heading:" in cell_lower:
                            heading = cell.split("heading:")[1].split("text:")[0].strip() if "text:" in cell_lower else cell.split("heading:")[1].strip()
                        if "text:" in cell_lower:
                            text = cell.split("text:")[1].strip() if "text:" in cell_lower else None

                        if heading or text:
                            if heading:
                                new_doc.add_paragraph(heading, style="Heading 1")
                            if text:
                                new_doc.add_paragraph(text, style="Normal")

                            # Create a new table with 8 columns
                            current_table = new_doc.add_table(rows=0, cols=8)
                            current_table.style = "Table Grid"
                            skip_row = True
                            break  # Stop processing this row further

                if skip_row:
                    continue  # Skip processing further for this row

                # Add data to table if it exists
                if current_table:
                    # Filter non-empty cells and limit to 8 columns
                    non_empty_row = [str(cell).strip() for cell in row if cell]  # Convert all valid cells to strings
                    if non_empty_row:
                        word_row = current_table.add_row().cells  # Get the new row

                        # Fill row cells (up to 8 columns)
                        for col_index, value in enumerate(non_empty_row[:8]):
                            cell = word_row[col_index]
                            run = cell.paragraphs[0].add_run(value)
                            if col_index == 0:  # Bold first column
                                run.bold = True
                            run.font.size = Pt(10.5)

                        # Merge columns 2-8 into one cell
                        merged_cell = word_row[1]
                        for col_index in range(2, 8):
                            merged_cell.merge(word_row[col_index])

        except Exception as e:
            print(f"Error processing {file_name}: {e}")

    # Save the Word document after processing all Excel files
    new_doc.save(OUTPUT_FILE)
    print(f"✅ Data saved in {OUTPUT_FILE}")

# Run the function
if __name__ == "__main__":
    PO_PSO()
