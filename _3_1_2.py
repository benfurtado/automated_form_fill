from docx import Document
import os
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_course_info(paragraphs):
    """Extracts course code and class/semester from document paragraphs."""
    course_info = {
        "course_code": "Not Found",
        "class_semester": "Not Found"
    }
    
    for para in paragraphs:
        text = para.text.strip()
        if "Course Code and Name:" in text:
            course_info["course_code"] = text.split(":", 1)[1].strip().replace("\u2013", "-")
        elif "Class and Semester:" in text:
            course_info["class_semester"] = text.split(":", 1)[1].strip()
    
    return course_info

def extract_revised_co_po_table(doc):
    """Extracts the table following the 'Revised CO-PO Mapping' heading."""
    table_data = []
    
    # Find the paragraph that mentions "Revised CO-PO Mapping"
    target_para_found = False
    for para in doc.paragraphs:
        if "revised co-po mapping" in para.text.lower():
            target_para_found = True
            break
    
    if not target_para_found:
        return table_data
    
    # Find the first table in the document
    if doc.tables:
        table = doc.tables[6]  # Assuming the 7th table is the required one
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if cells and any(x in cells[0].lower() for x in ["total", "average"]):
                continue
            if any(cell.strip() for cell in cells):
                table_data.append(cells)
    
    return table_data

def process_docx_files():
    """Processes all .docx files in 'temp_files' and creates a combined report."""
    output_file = "output.docx"
    output_doc = Document(output_file) if os.path.exists(output_file) else Document()
    
    temp_dir = "temp_files"
    files = sorted(
        [f for f in os.listdir(temp_dir) if f.endswith(".docx") and f.startswith("18")],
        key=lambda x: re.findall(r'\d{4}-\d{2}', x)
    )
    
    for filename in files:
        file_path = os.path.join(temp_dir, filename)
        doc = Document(file_path)
        
        # Extract course info
        course_info = extract_course_info(doc.paragraphs)
        
        # Extract CO-PO table
        co_po_table = extract_revised_co_po_table(doc)
        
        # Format filename with year
        year_match = re.findall(r'\d{4}-\d{2}', filename)
        year = year_match[0] if year_match else "Unknown"
        
        # Add a blank line for spacing
        output_doc.add_paragraph()
        
        # Add heading
        para = output_doc.add_paragraph(f"Class: {course_info['class_semester']} Subject: {course_info['course_code']}")
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(14)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if co_po_table:
            table = output_doc.add_table(rows=len(co_po_table)-1, cols=len(co_po_table[0]))
            table.style = "Table Grid"
            
            # Data rows (skip the first row)
            for row_idx, row in enumerate(co_po_table[1:]):
                new_row = table.rows[row_idx].cells
                for i, cell in enumerate(row):
                    new_row[i].text = cell
        else:
            output_doc.add_paragraph("Revised CO-PO Mapping table not found").italic = True

    # Save document
    output_doc.save(output_file)
    print(f"\nReport saved: {output_file}")
    
    # Open document
    open_document(output_file)

def open_document(file_path):
    """Opens a document in a platform-independent way."""
    import platform
    system = platform.system()
    
    if system == "Windows":
        os.startfile(file_path)
    elif system == "Darwin":  # macOS
        os.system(f"open \"{file_path}\"")
    elif system == "Linux":
        os.system(f"xdg-open \"{file_path}\"")

# Run the function
process_docx_files()
