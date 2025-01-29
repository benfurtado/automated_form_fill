from docx import Document
import os
from connection import drive, list_files_in_folder
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.shared import Pt
import merged

# Create a temporary directory for storing intermediate files
TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)

# Initialize a list to store all file titles
file_titles = []

# Function to adjust the margins of the document
def adjust_margins(doc):
    sections = doc.sections
    for section in sections:
        # Set the left margin (in inches)
        section.left_margin = Inches(0.5)  # Set to 0.5 inches or any other value you prefer
        section.right_margin = Inches(0.5)
        # You can also adjust other margins (top, right, bottom) if needed
        # section.top_margin = Inches(1)
        # section.right_margin = Inches(0.5)
        # section.bottom_margin = Inches(1)

# Function to process DOCX files
def process_docx_file(file):
    try:
        print(f"Processing file: {file['title']}")

        # Add the file title to the list
        file_titles.append(file['title'])

        # Download the file content as binary data to the temp folder
        temp_file_path = os.path.join(TEMP_DIR, file['title'])
        file.GetContentFile(temp_file_path)

        # Open the .docx file
        doc = Document(temp_file_path)

        # Initialize variables to store extracted data
        course_code = None
        subject = None
        average_row = None

        # Extract course code and subject
        for para in doc.paragraphs:
            if "Course Code and Name:" in para.text:
                line = para.text.split("Course Code and Name:")[1].strip()
                if "–" in line:
                    course_code, subject = line.split("\u2013", 1)
                    course_code = course_code.strip()
                    subject = subject.strip()
                break

        # Locate "Revised CO-PO Mapping:" and extract the "Average" row
        table_found = False
        for para_idx, para in enumerate(doc.paragraphs):
            if "Revised CO-PO Mapping:" in para.text:
                table_found = True
                print(f"Found 'Revised CO-PO Mapping:' at line {para_idx + 1}")

                # Locate the next table after this paragraph
                for table in doc.tables:
                    # Check if the table comes after the paragraph in the document flow
                    if para.text in table._element.xpath("./preceding-sibling::w:p//w:t/text()"):
                        print("Found table linked to 'Revised CO-PO Mapping:'")

                        # Extract the "Average" row
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells]
                            if "Average" in row_text:
                                # Exclude the "Average" column (first column)
                                average_row = row_text[1:]
                                break
                        break
                break

        if not table_found:
            print("No 'Revised CO-PO Mapping:' section found.")
        elif not average_row:
            print("No 'Average' row found in the table linked to 'Revised CO-PO Mapping:'.")

        return {"course_code": course_code, "subject": subject, "average_row": average_row}
    except Exception as e:
        print(f"Error processing DOCX file {file['title']}: {e}")
        return None

# Function to write extracted data into a properly formatted table
def write_data_to_table(extracted_data, doc):
    try:
        # Define headers for the table
        headers = [
            "Sr No", "Course Code", "Subject", "PO1", "PO2", "PO3", "PO4",
            "PO5", "PO6", "PO7", "PO8", "PO9", "PO10", "PO11", "PO12",
            "PSO1", "PSO2", "PSO3"
        ]

        # Define column widths in inches
        column_widths = [
            Inches(0.9),  # Sr No
            Inches(1.9),  # Course Code
            Inches(2.9),  # Subject
        ] + [Inches(0.9)] * (len(headers) - 3)  # Uniform width for PO/PSO columns

        # Create a new table or use the existing one
        if not doc.tables:
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            table.alignment = 0  # Align to left

            # Populate the header row
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header

        else:
            table = doc.tables[0]  # Use the first existing table

        # Adjust column widths
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(column_widths):
                    cell.width = column_widths[col_idx]

        # Populate table rows with extracted data
        start_sr_no = len(table.rows)  # Current number of rows, accounting for the header
        for idx, data in enumerate(extracted_data, start=start_sr_no):
            row = table.add_row().cells
            row[0].text = str(idx)  # Sr No
            row[1].text = data.get("course_code", " ")  # Course Code
            row[2].text = data.get("subject", " ")  # Subject

            # Fill PO1 to PSO3 values
            average_row = data.get("average_row", [])
            for col_idx, po_value in enumerate(average_row, start=3):
                if col_idx < len(headers):
                    row[col_idx].text = po_value or " "

        # Apply formatting to each cell
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)  # Set font size to 8 for a cleaner fit

        print("Table successfully updated with formatted widths and data.")
    except Exception as e:
        print(f"Error writing data to table: {e}")

# Function to process folders and files recursively
def process_folders(folder_id, doc):
    try:
        file_list = list_files_in_folder(folder_id)
        extracted_data = []

        for file in file_list:
            if file["mimeType"] == "application/vnd.google-apps.folder":  # Folder
                print(f"Entering folder: {file['title']}")
                process_folders(file["id"], doc)
            elif file["title"].startswith("18") and file["title"].endswith(".docx"):  # DOCX files starting with '18'
                data = process_docx_file(file)
                if data:
                    extracted_data.append(data)

        # Write data to the document after processing all files in this folder
        if extracted_data:
            write_data_to_table(extracted_data, doc)

    except Exception as e:
        print(f"Error processing folder ID {folder_id}: {e}")

# Main function
def main():
    try:
        root_folder_id = "1Bbp_TRb2dt-oRcKo3C7vHK7AHf0Hy90p"
        output_file_path = "output.docx"

        doc = Document('output.docx')
        adjust_margins(doc)  # Adjust margins to reduce the left margin

        merged.CO_Table()

        doc.add_heading("Extracted Information", 0)

        process_folders(root_folder_id, doc)
        doc.save(output_file_path)
        print(f"Data successfully saved in {output_file_path}")

        # Print all file titles
        print("\nProcessed File Titles:")
        for title in file_titles:
            print(title)

        # Auto-open the output file
        os.startfile(output_file_path)
    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()