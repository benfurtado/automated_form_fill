from docx import Document
import os
import platform
from connection import drive, list_files_in_folder  # Ensure connection.py is correctly set up
from docx.shared import Inches, Pt

# Create a temporary directory for storing intermediate files
TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)

file_titles = []  # Store processed file titles

# Function to adjust document margins
def adjust_margins(doc):
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

# Function to process DOCX files
def process_docx_file(file):
    try:
        print(f"Processing file: {file['title']}")
        file_titles.append(file['title'])

        # Download file to temporary storage
        temp_file_path = os.path.join(TEMP_DIR, file['title'])
        file.GetContentFile(temp_file_path)

        # Open .docx file
        doc = Document(temp_file_path)
        course_code, subject, average_row = None, None, None

        # Extract course code & subject
        for para in doc.paragraphs:
            if "Course Code and Name:" in para.text:
                line = para.text.split("Course Code and Name:")[1].strip()
                if "–" in line:
                    course_code, subject = line.split("\u2013", 1)
                    course_code, subject = course_code.strip(), subject.strip()
                break

        # Extract "Average" row from "Revised CO-PO Mapping" table
        for para in doc.paragraphs:
            if "Revised CO-PO Mapping:" in para.text:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        if "Average" in row_text:
                            average_row = row_text[1:]  # Exclude "Average" column
                            break
                break

        return {"course_code": course_code, "subject": subject, "average_row": average_row}
    except Exception as e:
        print(f"Error processing {file['title']}: {e}")
        return None

# Function to write extracted data into a table
def write_data_to_table(extracted_data, doc):
    try:
        headers = ["Sr No", "Course Code", "Subject"] + [f"PO{i}" for i in range(1, 13)] + ["PSO1", "PSO2", "PSO3"]
        column_widths = [Inches(0.9), Inches(1.9), Inches(2.9)] + [Inches(0.9)] * (len(headers) - 3)

        # Create table if not present
        table = doc.tables[0] if doc.tables else doc.add_table(rows=1, cols=len(headers))
        table.style, table.alignment = "Table Grid", 0

        # Fill header row
        for col_idx, header in enumerate(headers):
            table.cell(0, col_idx).text = header

        # Adjust column widths
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx < len(column_widths):
                    cell.width = column_widths[col_idx]

        # Fill table rows with extracted data
        for idx, data in enumerate(extracted_data, start=len(table.rows)):
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text = str(idx), data.get("course_code", ""), data.get("subject", "")
            for col_idx, po_value in enumerate(data.get("average_row", []), start=3):
                row[col_idx].text = po_value or " "

        # Set font size for better readability
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        print("Table updated successfully.")
    except Exception as e:
        print(f"Error writing table: {e}")

# Function to process Google Drive folders and files
def process_folders(folder_id, doc):
    try:
        file_list = list_files_in_folder(folder_id)
        extracted_data = []

        for file in file_list:
            if file["mimeType"] == "application/vnd.google-apps.folder":
                print(f"Entering folder: {file['title']}")
                process_folders(file["id"], doc)
            elif file["title"].startswith("18") and file["title"].endswith(".docx"):
                data = process_docx_file(file)
                if data:
                    extracted_data.append(data)

        if extracted_data:
            write_data_to_table(extracted_data, doc)
    except Exception as e:
        print(f"Error processing folder {folder_id}: {e}")

# Function to add a section break and adjust margins for a new page
def adjust_page_margins(doc):
    new_section = doc.add_paragraph().add_run()
    new_section.add_break()  # Section break (next page)
    doc.sections[-1].left_margin = Inches(0.5)
    doc.sections[-1].right_margin = Inches(0.5)

# Function to merge documents
def combined(doc):
    combined_path = "output.docx"
    combined_doc = Document(combined_path) if os.path.exists(combined_path) else Document()

    # Append paragraphs
    for para in doc.paragraphs:
        new_para = combined_doc.add_paragraph(para.text)
        new_para.style = "Heading 1"
        new_para.style.font.size, new_para.style.font.name = Pt(12), "Arial"

    # Append tables
    for table in doc.tables:
        combined_doc._element.append(table._element)

    adjust_page_margins(combined_doc)
    combined_doc.save(combined_path)
    print(f"Combined document saved as {combined_path}")

    open_document(combined_path)

# Function to open the output file (Cross-platform)
def open_document(file_path):
    system = platform.system()
    if system == "Windows":
        os.startfile(file_path)
    elif system == "Darwin":  # macOS
        os.system(f"open \"{file_path}\"")
    elif system == "Linux":
        os.system(f"xdg-open \"{file_path}\"")

# Main function
def main():
    try:
        root_folder_id = "1Bbp_TRb2dt-oRcKo3C7vHK7AHf0Hy90p"
        file_path = "temp.docx"

        doc = Document()
        adjust_margins(doc)
        process_folders(root_folder_id, doc)
        doc.save(file_path)
        print(f"Data saved in {file_path}")
        combined(doc)
    except Exception as e:
        print(f"An error occurred: {e}")

# Run the script
if __name__ == "__main__":
    main()
