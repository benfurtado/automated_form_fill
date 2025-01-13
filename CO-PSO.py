import subprocess
from docx import Document

# Open the source Word document
source_file_path = '22 co-po mapping corelation matrix.docx'  # Update with your file name
doc = Document(source_file_path)

# Create a new Word document to save content
new_doc = Document()

# Extract and copy text content
for paragraph in doc.paragraphs:
    if paragraph.text.strip():  # Ignore empty paragraphs
        new_doc.add_paragraph(paragraph.text, style=paragraph.style)

# Extract and process table content
for table in doc.tables:
    if len(table.rows) == 0:  # Skip empty tables
        continue

    # Identify the number of columns (excluding unwanted columns 14, 15, 16)
    original_columns = len(table.rows[0].cells)
    valid_columns = [idx for idx in range(original_columns) if idx  in {14, 15, 16}]
    
    # Create a new table in the output document with the reduced column count
    new_table = new_doc.add_table(rows=0, cols=len(valid_columns))
    new_table.style = 'Table Grid'

    # Iterate through rows, skipping the first one
    for row_index, row in enumerate(table.rows):
        if row_index == 0:  # Skip the first row
            continue

        # Add a new row to the output table
        new_row = new_table.add_row().cells

        # Add only the valid columns
        for new_idx, original_idx in enumerate(valid_columns):
            new_row[new_idx].text = row.cells[original_idx].text

# Save the new document
output_file_path = 'Filtered_Content.docx'  # Update with your desired output file name
new_doc.save(output_file_path)

# Automatically open the saved Word document
print(f"Content saved to {output_file_path}. Opening the file...")
subprocess.Popen(['open', output_file_path])

