import os
from docx import Document

# Open the source Word document
source_file_path = '22 co-po mapping corelation matrix.docx'  # Update with your file name
doc = Document(source_file_path)

# Create a new Word document to save the extracted rows
new_doc = Document()

# Extract second row as header and the average row from tables
for table in doc.tables:
    if len(table.rows) < 2:  # Skip tables with fewer than 2 rows
        continue

    # Get the second row as the header (omit the topmost row)
    header_row = table.rows[1]

    # Get the average row (last row)
    average_row = table.rows[-1]

    # Create a new table with the same number of columns as the source table
    num_columns = len(header_row.cells)
    new_table = new_doc.add_table(rows=0, cols=num_columns)
    new_table.style = 'Table Grid'

    # Add the header row (second row from the source table) to the new table
    header_cells = new_table.add_row().cells
    for col_idx, cell in enumerate(header_row.cells):
        header_cells[col_idx].text = cell.text

    # Add the average row (last row from the source table) to the new table
    average_cells = new_table.add_row().cells
    for col_idx, cell in enumerate(average_row.cells):
        average_cells[col_idx].text = cell.text

# Save the new document
output_file_path = 'Filtered_SecondRow_Average.docx'  # Desired output file name
new_doc.save(output_file_path)

# Automatically open the saved Word document
print(f"Filtered content saved to {output_file_path}. Opening the file...")
os.startfile(output_file_path)
