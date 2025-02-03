from docx import Document
from docx.shared import Pt
import os
import _3_1
import _3_1_1
import _3_1_2
import _3_1_3


# Function to add bold text with line spacing
def add_text_to_doc(file, text):
    doc = Document(file)

    # Add formatted text
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)

    # Set paragraph spacing
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    doc.save(file)


# File name
file = "output.docx"

# Create a new document if it doesn't exist
if not os.path.exists(file):
    Document().save(file)

# Add text and call functions
add_text_to_doc(file, "3.1 Establish the correlation between the courses and the Program Outcomes (POs) and Program Specific Outcomes (PSOs) (20)\n(Program Outcomes as mentioned in Annexure I and Program Specific Outcomes as defined by the Program)")
_3_1.PO_PSO()

add_text_to_doc(file, "3.1.1 Course Outcomes (COs) (SAR should include course outcomes of one course from each semester of study, however, should be prepared for all courses and made available as evidence, if asked) (05)\nNote: Number of Outcomes for a Course is expected to be around 6.")
_3_1_1.CO_Table()

add_text_to_doc(file, "3.1.2 CO-PO matrices of courses selected in 3.1.1 (six matrices to be mentioned; one per semester from 3rd to 8th semester) (05)")
_3_1_2.process_docx_files()

add_text_to_doc(file, "3.1.3 Program level Course-PO matrix of all courses INCLUDING first year courses (10)")
_3_1_3.main()

print(f"Data successfully saved in {file}")

# Open the file in a platform-independent way
import platform
if platform.system() == "Windows":
    os.startfile(file)
elif platform.system() == "Darwin":  # macOS
    os.system(f"open {file}")
elif platform.system() == "Linux":
    os.system(f"xdg-open {file}")
